import io
import docx
from typing import List, Dict, Any, Union
from app.rag.parsers.base import BaseParser


class DocxParser(BaseParser):
    """
    Parser for extracting text from Microsoft Word (.docx) files.
    Groups paragraphs into logical page boundaries to support chunk indexing.
    """
    def parse(self, file_input: Union[bytes, str]) -> List[Dict[str, Any]]:
        if isinstance(file_input, str):
            doc = docx.Document(file_input)
        else:
            doc = docx.Document(io.BytesIO(file_input))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        pages = []
        paragraphs_per_page = 15
        current_page_paragraphs = []
        page_num = 1

        for idx, text in enumerate(paragraphs):
            current_page_paragraphs.append(text)
            if (idx + 1) % paragraphs_per_page == 0:
                pages.append({
                    "text": "\n\n".join(current_page_paragraphs),
                    "page_number": page_num
                })
                current_page_paragraphs = []
                page_num += 1

        # Append remaining paragraphs
        if current_page_paragraphs:
            pages.append({
                "text": "\n\n".join(current_page_paragraphs),
                "page_number": page_num
            })

        return pages
